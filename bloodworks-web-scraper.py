# Python script for web scraping to extract data from a website
import requests, sys, datetime
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions
import csv, docx, traceback
from docx import Document
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import StaleElementReferenceException

aos = [
    "Alpines",
    "Angmar",
    "Bloodsport",
    "Bobcat",
    "Bonsai",
    "Columbia City",
    "The Combine",
    "Counterbalance",
    "Doom",
    "Flash",
    "Gasworks",
    "Grasslawn",
    "Hawks Nest",
    "Heritage",
    "Hiawatha",
    "Hurricane Ridge",
    "KiMS",
    "Log Boom",
    "Mohai",
    "PBJ - City Hall",
    "PBJ - Doom",
    "PBJ - Google Monday",
    "PBJ - Google Thursday",
    "PBJ - Ravenna",
    "Purple Haze",
    "Perestroika",
    "Rat City",
    "Robinswood",
    "Ruck Mountain",
    "Sasquatch",
    "Soft Trot",
    "Space Needle",
    "Speakeasy",
    "Thunderdome",
    "Timber",
    "Torque",
    "Tundra",
    "Tundra Speed",
    "Valhalla"]

zipcodes = [
    98027,
    98021,
    98036,
    98052,
    98112,
    98118,
    98117,
    98028,
    98115,
    98103,
    98043,
    98033,
    98116,
    98029,
    98109,
    98105,
    98005,
    98146,
    98007,
    98034,
    98056,
    98072            
]

zipcodesandaos = {
    98005: "Perestroika",
    98007: "Robinswood",
    98021: "Angmar",
    98027: "Alpines",
    98028: "Doom, Log Boom, PBJ - Doom, Soft Trot",
    98029: "Hurricane Ridge",
    98033: "Heritage, KiMS, PBJ - City Hall, PBJ - Google Mondaya, PBJ - Google Thursday, Ruck Mountain, Torque",
    98034: "Sasquatch, Speakeasy, Thunderdome",
    98036: "Bloodsport",
    98043: "Hawks Nest",
    98052: "Bobcat, Grasslawn",
    98056: "Purple Haze",
    98072: "Timber, Tundra, Tundra Speed",
    98103: "Gasworks",
    98105: "PBJ - Ravenna",
    98109: "Mohai, Space Needle",
    98112: "Bonsai",
    98115: "Flash, Valhalla",
    98116: "Hiawatha",
    98117: "The Combine",
    98118: "Columbia City, Counterbalance",
    98146: "Rat City"
}

aosandzipcodes = {
    "Alpines": 98027,
    "Angmar": 98021,
    "Bloodsport": 98036,
    "Bobcat": 98052,
    "Bonsai": 98112,
    "Columbia City": 98118,
    "The Combine": 98117,
    "Counterbalance": 98118,
    "Doom": 98028,
    "Flash": 98115,
    "Gasworks": 98103,
    "Grasslawn": 98052,
    "Hawks Nest": 98043,
    "Heritage": 98033,
    "Hiawatha": 98116,
    "Hurricane Ridge": 98029,
    "KiMS": 98033,
    "Log Boom": 98028,
    "Mohai": 98109,
    "PBJ - City Hall": 98033,
    "PBJ - Doom": 98028,
    "PBJ - Google Monday": 98033,
    "PBJ - Google Thursday": 98033,
    "PBJ - Ravenna": 98105,
    "Perestroika": 98005,
    "Purple Haze": 98056,
    "Rat City": 98146,
    "Robinswood": 98007,
    "Ruck Mountain": 98033,
    "Sasquatch": 98034,
    "Soft Trot": 98028,
    "Space Needle": 98109,
    "Speakeasy": 98034,
    "Thunderdome": 98034,
    "Timber": 98072,
    "Torque": 98033,
    "Tundra": 98072,
    "Tundra Speed": 98072,
    "Valhalla": 98115
}


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    # new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink

# url https://donate.bloodworksnw.org/donor/schedules/zip

url = ""
today = datetime.date.today()
endDate = today + datetime.timedelta(days=30)
print('endDate: ', endDate)
datestring = str(endDate.month) + "/" + str(endDate.day) + "/" + str(endDate.year)
print('datestring: ', datestring)
url = "https://donate.bloodworksnw.org/donor/schedules/zip"

# if (len(sys.argv) == 0):
#     url = "https://donate.bloodworksnw.org/donor/schedules/zip"
# else:
#     url = sys.argv[1]
if (len(sys.argv) > 1):
    datestring = sys.argv[1]

response = requests.get(url)
soup = BeautifulSoup(response.text, 'html.parser')
picks = soup.find_all('p' and 'strong')

    
# Your code here to extract relevant data from the website

# automation.py


chrome_driver_path = './chromedriver'

service = Service(chrome_driver_path)
driver = webdriver.Chrome()

driver.get(url)
wait = WebDriverWait(driver, timeout=2)
zipcode_box = driver.find_element(By.NAME, "zipcode")
wait.until(lambda d : zipcode_box.is_displayed())
distanceSlider = driver.find_element(By.CLASS_NAME, "slider-handle")
i = 22
while i > 0:
    distanceSlider.send_keys(Keys.ARROW_LEFT)
    i = i - 1

document = Document()
for z in zipcodes:
    header = zipcodesandaos[z]
    
    print('header: ', header, '; z: ', z)
    try: 
        zipcode_box = driver.find_element(By.NAME, "zipcode")
        zipcode_box.click()
        zipcode_box.send_keys(Keys.HOME)
        zipcode_box.send_keys(z)
        searchBtn = driver.find_element(By.ID, "search")

        endDateBox = driver.find_element(By.NAME, "end_date")
        endDateBox.click()
        # endDateBox.send_keys(Keys.END)
        # endDateBox.send_keys(Keys.SHIFT + Keys.HOME)
        # endDateBox.send_keys(Keys.DELETE)
        for letter in datestring:
            endDateBox.send_keys(letter)
        endDateBox.send_keys(Keys.TAB)
        searchBtn.click()

        driver.implicitly_wait(10)
        ignored_exceptions=(NoSuchElementException,StaleElementReferenceException,)
        WebDriverWait(driver=driver, timeout=30,ignored_exceptions=ignored_exceptions).until(expected_conditions.presence_of_element_located((By.ID, "item_table")))
        resultsTable = driver.find_element(By.ID, "item_table")
        wait.until(lambda d : resultsTable.is_displayed())

        WebDriverWait(driver=driver, timeout=30,ignored_exceptions=ignored_exceptions).until(expected_conditions.presence_of_element_located((By.XPATH, "//*[@id=\"item_table\"]/thead/tr/th[4]")))
        dateCol = resultsTable.find_element(By.XPATH, "//*[@id=\"item_table\"]/thead/tr/th[4]") #//*[@id="item_table"]/thead/tr/th[4]
        wait.until(lambda d : dateCol.is_displayed())
        dateCol.click()
        WebDriverWait(driver=driver, timeout=30,ignored_exceptions=ignored_exceptions).until(expected_conditions.presence_of_element_located((By.TAG_NAME, "tr")))
        results = resultsTable.find_elements(By.TAG_NAME, "tr")
        print(len(results))
        filename = datestring.replace("/", "_")
        fields = ['AOs', 'Drive Name', 'Drive Date/Time', 'Drive Link']
        rows = []
        paginateBtns = len(driver.find_elements(By.CLASS_NAME, "paginate_button")) - 2
        p = 0
        while p < paginateBtns:
            WebDriverWait(driver=driver, timeout=30,ignored_exceptions=ignored_exceptions).until(expected_conditions.presence_of_element_located((By.ID, "item_table")))
            resultsTable = driver.find_element(By.ID, "item_table")
            wait.until(lambda d : resultsTable.is_displayed())

            results = resultsTable.find_elements(By.TAG_NAME, "tr")
            nextButton = driver.find_element(By.ID, "item_table_next")
            nextButtonATag = nextButton.find_element(By.TAG_NAME, "a")
            for row in results:
                cols = row.find_elements(By.TAG_NAME, "td")
                if len(cols) > 0:
                    driveName = cols[0].text
                    driveDateTime = cols[3].text
                    driveDateTime = driveDateTime.replace('\n', ' ')
                    driveLink = cols[4].find_element(By.TAG_NAME, 'a').get_attribute("href")
                    rows.append([zipcodesandaos[z], driveName, driveDateTime, driveLink])
            nextButtonATag.click()
            p = p + 1

        blooddrivescsv = open(filename + ".csv", 'w+')
        with blooddrivescsv as csvfile:
            
            # creating a csv writer object  
            csvwriter = csv.writer(csvfile)  
                
            # writing the fields  
            csvwriter.writerow(fields)  
                
            # writing the data rows  
            csvwriter.writerows(rows) 

        document.add_heading(str(header) + ' Blood Drives through ' + datestring)
        for row in rows:
            p = document.add_paragraph('')
            add_hyperlink(p, row[1] + ", " + row[2], row[3] )
    except Exception as e:
        print('e: ', e )
        print(traceback.print_exc())

document.save("blood-drives-" + filename + ".docx")

driver.quit()